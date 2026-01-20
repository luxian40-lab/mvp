"""
Utilidades para generación de cursos con IA
"""

import os
import json
import logging
from io import BytesIO
from typing import Dict, List, Optional

import openai
from PyPDF2 import PdfReader
from docx import Document

from .models import Curso, Modulo, PreguntaModulo, Examen, PreguntaExamen, Cliente

logger = logging.getLogger(__name__)

# Mapeo de modelos disponibles
MODELOS_IA = {
    'gpt-3.5-turbo': {'proveedor': 'openai', 'nombre': 'GPT-3.5 Turbo', 'costo': 0.01},
    'gpt-4': {'proveedor': 'openai', 'nombre': 'GPT-4', 'costo': 0.30},
    'claude-3-sonnet': {'proveedor': 'anthropic', 'nombre': 'Claude 3 Sonnet', 'costo': 0.015},
    'claude-3-opus': {'proveedor': 'anthropic', 'nombre': 'Claude 3 Opus', 'costo': 0.40},
    'gemini-pro': {'proveedor': 'google', 'nombre': 'Gemini Pro', 'costo': 0.005},
}


def extraer_texto_pdf(archivo) -> str:
    """
    Extrae texto de un archivo PDF.
    
    Args:
        archivo: Archivo PDF (InMemoryUploadedFile o similar)
        
    Returns:
        str: Texto extraído del PDF
    """
    try:
        pdf_reader = PdfReader(archivo)
        texto = ""
        
        for pagina in pdf_reader.pages:
            texto += pagina.extract_text() + "\n\n"
        
        logger.info(f"PDF procesado: {len(pdf_reader.pages)} páginas, {len(texto)} caracteres")
        return texto.strip()
    
    except Exception as e:
        logger.error(f"Error al extraer texto de PDF: {e}")
        raise ValueError(f"No se pudo procesar el PDF: {str(e)}")


def extraer_texto_word(archivo) -> str:
    """
    Extrae texto de un archivo Word (.docx).
    
    Args:
        archivo: Archivo Word (InMemoryUploadedFile o similar)
        
    Returns:
        str: Texto extraído del documento
    """
    try:
        doc = Document(archivo)
        texto = ""
        
        for parrafo in doc.paragraphs:
            if parrafo.text.strip():
                texto += parrafo.text + "\n\n"
        
        logger.info(f"Word procesado: {len(doc.paragraphs)} párrafos, {len(texto)} caracteres")
        return texto.strip()
    
    except Exception as e:
        logger.error(f"Error al extraer texto de Word: {e}")
        raise ValueError(f"No se pudo procesar el documento Word: {str(e)}")


def extraer_texto_txt(archivo) -> str:
    """
    Extrae texto de un archivo de texto plano.
    
    Args:
        archivo: Archivo de texto
        
    Returns:
        str: Texto del archivo
    """
    try:
        # Intentar diferentes encodings
        encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
        
        for encoding in encodings:
            try:
                archivo.seek(0)
                texto = archivo.read().decode(encoding)
                logger.info(f"TXT procesado con encoding {encoding}: {len(texto)} caracteres")
                return texto.strip()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("No se pudo decodificar el archivo de texto")
    
    except Exception as e:
        logger.error(f"Error al extraer texto de TXT: {e}")
        raise ValueError(f"No se pudo procesar el archivo de texto: {str(e)}")


def extraer_texto_documento(archivo) -> str:
    """
    Extrae texto de un documento (PDF, Word o TXT).
    
    Args:
        archivo: Archivo subido
        
    Returns:
        str: Texto extraído
    """
    nombre_archivo = archivo.name.lower()
    
    if nombre_archivo.endswith('.pdf'):
        return extraer_texto_pdf(archivo)
    elif nombre_archivo.endswith('.docx'):
        return extraer_texto_word(archivo)
    elif nombre_archivo.endswith('.txt'):
        return extraer_texto_txt(archivo)
    else:
        raise ValueError("Formato no soportado. Use PDF, Word (.docx) o TXT")


def generar_estructura_curso_con_ia(texto: str, modelo: str = "gpt-3.5-turbo") -> Dict:
    """
    Genera estructura de curso usando IA (OpenAI, Claude o Gemini).
    
    Args:
        texto: Contenido del curso extraído del documento
        modelo: Modelo de IA (gpt-3.5-turbo, gpt-4, claude-3-sonnet, gemini-pro, etc.)
        
    Returns:
        Dict: Estructura del curso en formato JSON
    """
    # Limitar texto a 12000 caracteres para no exceder límites de tokens
    texto_resumido = texto[:12000] if len(texto) > 12000 else texto
    
    # Detectar proveedor del modelo
    info_modelo = MODELOS_IA.get(modelo, {'proveedor': 'openai'})
    proveedor = info_modelo['proveedor']
    
    prompt = f"""
Eres un experto en diseño instruccional. Analiza el siguiente contenido de curso y genera una estructura de aprendizaje completa.

CONTENIDO DEL CURSO:
{texto_resumido}

Genera un JSON con esta estructura EXACTA (muy importante que sea JSON válido):

{{
  "titulo": "Título del curso (corto y descriptivo)",
  "descripcion": "Descripción de 2-3 líneas explicando qué aprenderán los estudiantes",
  "duracion_estimada": "Ejemplo: 4 semanas, 8 horas, etc.",
  "nivel": "Básico, Intermedio o Avanzado",
  "puntos_por_leccion": 50,
  "puntos_por_quiz": 100,
  "modulos": [
    {{
      "nombre": "Módulo 1: Introducción",
      "descripcion": "Breve descripción del módulo",
      "orden": 1,
      "lecciones": [
        {{
          "titulo": "Lección 1.1: Conceptos básicos",
          "contenido": "Resumen del contenido de la lección (3-5 párrafos)",
          "orden": 1,
          "duracion_minutos": 15,
          "preguntas": [
            {{
              "texto": "¿Pregunta de opción múltiple relacionada?",
              "opciones": [
                {{"texto": "Opción A (incorrecta)", "es_correcta": false}},
                {{"texto": "Opción B (correcta)", "es_correcta": true}},
                {{"texto": "Opción C (incorrecta)", "es_correcta": false}},
                {{"texto": "Opción D (incorrecta)", "es_correcta": false}}
              ],
              "explicacion": "Por qué la opción B es correcta"
            }}
          ]
        }}
      ],
      "mini_examen": [
        {{
          "texto": "Pregunta evaluativa del módulo",
          "opciones": [
            {{"texto": "Opción 1", "es_correcta": false}},
            {{"texto": "Opción 2", "es_correcta": true}},
            {{"texto": "Opción 3", "es_correcta": false}},
            {{"texto": "Opción 4", "es_correcta": false}}
          ],
          "explicacion": "Explicación de la respuesta correcta"
        }}
      ]
    }}
  ],
  "sugerencias_gamificacion": [
    "Ejemplo: Otorgar insignia 'Novato' al completar primer módulo",
    "Ejemplo: Desbloquear contenido bonus con 500 puntos"
  ],
  "temas_campanas": [
    "Recordatorio: Módulo 1 disponible",
    "Motivación: ¡Ya completaste el 50%!"
  ]
}}

INSTRUCCIONES IMPORTANTES:
1. Genera 3-5 módulos con 3-4 lecciones cada uno
2. Cada lección debe tener 2-3 preguntas
3. Cada módulo debe tener 5 preguntas en mini_examen
4. Las preguntas deben ser claras y relacionadas con el contenido
5. Solo UNA opción debe ser correcta (es_correcta: true)
6. El contenido de las lecciones debe ser educativo y bien estructurado
7. Responde SOLO con el JSON, sin texto adicional
"""

    try:
        logger.info(f"Enviando {len(texto_resumido)} caracteres a {modelo} ({proveedor})")
        
        # Llamar al proveedor correspondiente
        if proveedor == 'openai':
            contenido_respuesta = _generar_con_openai(modelo, prompt)
        elif proveedor == 'anthropic':
            contenido_respuesta = _generar_con_claude(modelo, prompt)
        elif proveedor == 'google':
            contenido_respuesta = _generar_con_gemini(modelo, prompt)
        else:
            raise ValueError(f"Proveedor de IA no soportado: {proveedor}")
        
        # Limpiar respuesta si viene con markdown
        if contenido_respuesta.startswith('```json'):
            contenido_respuesta = contenido_respuesta.replace('```json', '').replace('```', '').strip()
        elif contenido_respuesta.startswith('```'):
            contenido_respuesta = contenido_respuesta.replace('```', '').strip()
        
        estructura = json.loads(contenido_respuesta)
        
        logger.info(f"Estructura generada: {len(estructura.get('modulos', []))} módulos")
        
        return estructura
    
    except json.JSONDecodeError as e:
        logger.error(f"Error al parsear JSON de OpenAI: {e}")
        logger.error(f"Respuesta recibida: {contenido_respuesta[:500]}")
        raise ValueError("La IA generó una respuesta inválida. Intente nuevamente.")
    
    except Exception as e:
        logger.error(f"Error al generar estructura con IA: {e}")
        raise ValueError(f"Error al procesar con IA: {str(e)}")


def guardar_curso_desde_estructura(estructura: Dict, cliente: Cliente, archivo_nombre: str) -> Curso:
    """
    Guarda un curso en la base de datos desde la estructura generada por IA.
    
    Args:
        estructura: Estructura del curso generada por IA
        cliente: Cliente al que pertenece el curso
        archivo_nombre: Nombre del archivo original
        
    Returns:
        Curso: Instancia del curso creado
    """
    try:
        # Crear curso
        curso = Curso.objects.create(
            cliente=cliente,
            nombre=estructura['titulo'],
            descripcion=estructura['descripcion'],
            duracion_estimada=estructura.get('duracion_estimada', '4 semanas'),
            nivel=estructura.get('nivel', 'intermedio'),
            puntos_recompensa=estructura.get('puntos_por_leccion', 50),
            activo=False  # Inactivo hasta que admin revise
        )
        
        logger.info(f"Curso creado: {curso.nombre} (ID: {curso.id})")
        
        # Crear módulos
        for idx, modulo_data in enumerate(estructura['modulos'], 1):
            modulo = Modulo.objects.create(
                curso=curso,
                numero=idx,
                titulo=modulo_data['nombre'],
                descripcion=modulo_data.get('descripcion', ''),
                contenido=modulo_data.get('contenido', modulo_data.get('descripcion', '')),
                duracion_dias=7
            )
            
            logger.info(f"  Módulo creado: {modulo.titulo}")
            
            # Crear preguntas del mini-examen del módulo
            for pregunta_data in modulo_data.get('mini_examen', [])[:1]:  # Solo 1 pregunta por módulo
                # Encontrar la opción correcta
                opciones = pregunta_data.get('opciones', [])
                respuesta_correcta = 'A'
                opcion_a = opciones[0]['texto'] if len(opciones) > 0 else ''
                opcion_b = opciones[1]['texto'] if len(opciones) > 1 else ''
                opcion_c = opciones[2]['texto'] if len(opciones) > 2 else ''
                opcion_d = opciones[3]['texto'] if len(opciones) > 3 else ''
                
                for idx_opt, opcion in enumerate(opciones):
                    if opcion.get('es_correcta'):
                        respuesta_correcta = ['A', 'B', 'C', 'D'][idx_opt]
                        break
                
                PreguntaModulo.objects.create(
                    modulo=modulo,
                    pregunta=pregunta_data['texto'],
                    opcion_a=opcion_a,
                    opcion_b=opcion_b,
                    opcion_c=opcion_c or '',
                    opcion_d=opcion_d or '',
                    respuesta_correcta=respuesta_correcta,
                    explicacion=pregunta_data.get('explicacion', ''),
                    activa=True
                )
                
                logger.info(f"    Mini-examen agregado al módulo")
        
        # Crear examen final del curso (opcional)
        try:
            examen = Examen.objects.create(
                curso=curso,
                instrucciones="Responde las siguientes preguntas sobre el curso:",
                puntaje_minimo=70
            )
            logger.info(f"  Examen final creado")
        except Exception as e:
            logger.warning(f"No se pudo crear examen final: {e}")
        
        logger.info(f"✅ Curso completo guardado: {curso.nombre}")
        return curso
    
    except Exception as e:
        logger.error(f"Error al guardar curso: {e}")
        import traceback
        traceback.print_exc()
        # Si algo falla, eliminar el curso para evitar datos incompletos
        if 'curso' in locals():
            curso.delete()
        raise ValueError(f"Error al guardar el curso: {str(e)}")


def regenerar_seccion(curso_id: int, tipo: str, seccion_id: int, instrucciones: str = "") -> bool:
    """
    Regenera una sección específica del curso (módulo, lección, etc.).
    
    Args:
        curso_id: ID del curso
        tipo: Tipo de sección ('modulo', 'leccion', 'preguntas')
        seccion_id: ID de la sección a regenerar
        instrucciones: Instrucciones adicionales para la IA
        
    Returns:
        bool: True si se regeneró exitosamente
    """
    # TODO: Implementar en Fase 3
    logger.info(f"Regeneración de {tipo} {seccion_id} del curso {curso_id}")
    return True


def validar_estructura_curso(estructura: Dict) -> tuple[bool, List[str]]:
    """
    Valida que la estructura generada tenga todos los campos necesarios.
    
    Args:
        estructura: Estructura del curso a validar
        
    Returns:
        tuple: (es_valida, lista_de_errores)
    """
    errores = []
    
    # Validar campos principales
    if 'titulo' not in estructura or not estructura['titulo']:
        errores.append("Falta el título del curso")
    
    if 'descripcion' not in estructura or not estructura['descripcion']:
        errores.append("Falta la descripción del curso")
    
    if 'modulos' not in estructura or not estructura['modulos']:
        errores.append("El curso debe tener al menos un módulo")
    else:
        # Validar módulos
        for idx, modulo in enumerate(estructura['modulos']):
            if 'nombre' not in modulo:
                errores.append(f"Módulo {idx+1}: falta nombre")
            
            if 'lecciones' not in modulo or not modulo['lecciones']:
                errores.append(f"Módulo {idx+1}: debe tener al menos una lección")
            else:
                # Validar lecciones
                for idx_l, leccion in enumerate(modulo['lecciones']):
                    if 'titulo' not in leccion:
                        errores.append(f"Módulo {idx+1}, Lección {idx_l+1}: falta título")
                    
                    if 'contenido' not in leccion:
                        errores.append(f"Módulo {idx+1}, Lección {idx_l+1}: falta contenido")
    
    return (len(errores) == 0, errores)


# ========================================
# FUNCIONES AUXILIARES POR PROVEEDOR DE IA
# ========================================

def _generar_con_openai(modelo: str, prompt: str) -> str:
    """Genera contenido usando OpenAI (GPT-3.5 o GPT-4)"""
    try:
        client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        
        response = client.chat.completions.create(
            model=modelo,
            messages=[
                {"role": "system", "content": "Eres un experto en diseño instruccional. Respondes SOLO con JSON válido."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        
        return response.choices[0].message.content.strip()
    
    except Exception as e:
        logger.error(f"Error con OpenAI: {e}")
        raise ValueError(f"Error al comunicarse con OpenAI: {str(e)}")


def _generar_con_claude(modelo: str, prompt: str) -> str:
    """Genera contenido usando Anthropic Claude"""
    try:
        # Intentar importar anthropic
        try:
            import anthropic
        except ImportError:
            raise ValueError("Para usar Claude, instala: pip install anthropic")
        
        client = anthropic.Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))
        
        message = client.messages.create(
            model=modelo,
            max_tokens=4000,
            temperature=0.7,
            system="Eres un experto en diseño instruccional. Respondes SOLO con JSON válido.",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        return message.content[0].text.strip()
    
    except Exception as e:
        logger.error(f"Error con Claude: {e}")
        raise ValueError(f"Error al comunicarse con Claude: {str(e)}")


def _generar_con_gemini(modelo: str, prompt: str) -> str:
    """Genera contenido usando Google Gemini"""
    try:
        # Intentar importar google-generativeai
        try:
            import google.generativeai as genai
        except ImportError:
            raise ValueError("Para usar Gemini, instala: pip install google-generativeai")
        
        genai.configure(api_key=os.getenv('GOOGLE_API_KEY'))
        
        model = genai.GenerativeModel(
            model_name='gemini-pro',
            generation_config={
                'temperature': 0.7,
                'max_output_tokens': 4000,
            }
        )
        
        response = model.generate_content(
            f"Eres un experto en diseño instruccional. Respondes SOLO con JSON válido.\n\n{prompt}"
        )
        
        return response.text.strip()
    
    except Exception as e:
        logger.error(f"Error con Gemini: {e}")
        raise ValueError(f"Error al comunicarse con Gemini: {str(e)}")
